from docx import Document
import os

def create_docx_file(filename, content):
    doc = Document()
    doc.add_heading(filename.replace('.docx', ''), level=1)
    
    for section in content:
        doc.add_heading(section['title'], level=2)
        doc.add_paragraph(section['content'])
    
    doc.save(os.path.join('documents', filename))
    print(f"Created: {filename}")

nlp_intro = [
    {
        'title': '什么是自然语言处理',
        'content': '自然语言处理（Natural Language Processing，简称NLP）是人工智能领域的一个重要分支，它致力于使计算机能够理解、解释和生成人类语言。NLP结合了计算机科学、语言学和人工智能技术，旨在实现人与计算机之间的自然语言交互。'
    },
    {
        'title': 'NLP的主要应用',
        'content': 'NLP技术在多个领域有着广泛的应用，包括：机器翻译、情感分析、文本分类、信息抽取、问答系统、语音识别、文本生成等。这些应用已经深入到我们日常生活的方方面面，如智能客服、语音助手、搜索引擎等。'
    },
    {
        'title': 'NLP的发展历程',
        'content': 'NLP的发展经历了几个重要阶段：早期基于规则的方法、统计机器学习方法，以及近年来兴起的深度学习方法。特别是Transformer架构的出现，极大地推动了NLP技术的发展，使得机器翻译、文本生成等任务取得了突破性进展。'
    }
]

transformer = [
    {
        'title': 'Transformer架构简介',
        'content': 'Transformer是一种基于自注意力机制的深度学习模型架构，由Google在2017年的论文《Attention is All You Need》中提出。它彻底改变了NLP领域，成为许多先进模型的基础，如BERT、GPT等。'
    },
    {
        'title': '自注意力机制',
        'content': '自注意力机制是Transformer的核心，它允许模型在处理序列数据时，同时关注序列中的所有位置。这种机制能够捕捉长距离依赖关系，相比传统的循环神经网络具有显著优势。'
    },
    {
        'title': 'Transformer的组成',
        'content': 'Transformer主要由编码器和解码器两部分组成。编码器负责处理输入序列，解码器负责生成输出序列。每一部分都包含多层堆叠的注意力层和前馈神经网络层。'
    }
]

bert = [
    {
        'title': 'BERT模型',
        'content': 'BERT（Bidirectional Encoder Representations from Transformers）是Google在2018年推出的预训练语言模型。它采用双向Transformer编码器，能够捕捉上下文信息，在多种NLP任务上取得了SOTA成绩。'
    },
    {
        'title': 'BERT的预训练任务',
        'content': 'BERT通过两个预训练任务进行训练：掩码语言模型（Masked Language Model，MLM）和下一句预测（Next Sentence Prediction，NSP）。MLM随机掩盖输入中的部分token，让模型预测被掩盖的token；NSP则让模型判断两个句子是否连续。'
    },
    {
        'title': 'BERT的应用',
        'content': 'BERT在多种NLP任务中表现出色，包括问答系统、文本分类、命名实体识别、情感分析等。通过微调，BERT可以适应各种特定任务的需求。'
    }
]

gpt = [
    {
        'title': 'GPT系列模型',
        'content': 'GPT（Generative Pre-trained Transformer）是OpenAI推出的一系列生成式语言模型。从GPT-1到GPT-4，模型规模和能力不断提升，展现了强大的文本生成能力。'
    },
    {
        'title': 'GPT的架构特点',
        'content': 'GPT采用仅解码器的Transformer架构，通过自回归方式生成文本。模型在大规模文本数据上进行预训练，学习语言的统计规律和知识。'
    },
    {
        'title': 'GPT的应用场景',
        'content': 'GPT在文本生成、对话系统、代码生成、摘要生成等多个领域有着广泛应用。它能够生成连贯、自然的文本，甚至可以完成复杂的推理任务。'
    }
]

rag = [
    {
        'title': '检索增强生成',
        'content': '检索增强生成（Retrieval-Augmented Generation，简称RAG）是一种结合信息检索和文本生成的技术。它通过从外部知识库中检索相关信息，为生成模型提供参考，从而提高回答的准确性和可靠性。'
    },
    {
        'title': 'RAG的工作流程',
        'content': 'RAG的典型工作流程包括：将文档分割成小块并进行向量化存储；当用户提出问题时，从向量数据库中检索最相关的文档片段；将检索到的文档作为上下文输入给语言模型，生成基于文档内容的回答。'
    },
    {
        'title': 'RAG的优势',
        'content': 'RAG技术具有以下优势：能够处理实时更新的知识；减少模型幻觉；提高回答的可解释性；支持特定领域的专业知识问答。这些优势使得RAG成为构建可靠问答系统的重要技术。'
    }
]

create_docx_file('自然语言处理简介.docx', nlp_intro)
create_docx_file('Transformer架构.docx', transformer)
create_docx_file('BERT模型详解.docx', bert)
create_docx_file('GPT系列模型.docx', gpt)
create_docx_file('检索增强生成RAG.docx', rag)

print("\n所有示例文档已创建完成！")